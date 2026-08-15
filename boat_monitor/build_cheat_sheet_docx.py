#!/usr/bin/env python3
"""Generate WIRING_CHEAT_SHEET.docx — printable bench/field wiring guide."""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

OUT = Path(__file__).parent / "WIRING_CHEAT_SHEET.docx"


def margins(doc):
    for s in doc.sections:
        s.top_margin = Inches(0.45)
        s.bottom_margin = Inches(0.45)
        s.left_margin = Inches(0.5)
        s.right_margin = Inches(0.5)


def h(doc, text, level=2):
    doc.add_heading(text, level=level)


def p(doc, text, bold=False, size=9):
    par = doc.add_paragraph()
    r = par.add_run(text)
    r.bold = bold
    r.font.size = Pt(size)


def bullets(doc, items, size=9):
    for item in items:
        par = doc.add_paragraph(style="List Bullet")
        r = par.add_run(item)
        r.font.size = Pt(size)


def tbl(doc, headers, rows, widths=None):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    for i, hd in enumerate(headers):
        c = t.rows[0].cells[i]
        c.text = hd
        for r in c.paragraphs[0].runs:
            r.bold = True
            r.font.size = Pt(8)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            c = t.rows[ri + 1].cells[ci]
            c.text = str(val)
            for r in c.paragraphs[0].runs:
                r.font.size = Pt(8)
    if widths:
        for row in t.rows:
            for i, w in enumerate(widths):
                row.cells[i].width = Inches(w)
    doc.add_paragraph()


def main():
    doc = Document()
    margins(doc)

    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("Boat Monitor P2 — Wiring Guide (print at bench)")
    r.bold = True
    r.font.size = Pt(14)
    s = doc.add_paragraph()
    s.alignment = WD_ALIGN_PARAGRAPH.CENTER
    s.add_run("HOUSE = middle INA260 · ENGINE = bottom INA260 · VSNS on Pico pin 15").font.size = Pt(9)

    h(doc, "Policy", 1)
    p(
        doc,
        "Section 1 is already wired — do not unsolder or move those connections. "
        "Everything else below is new pigtails, modem, opto, PlusRoc, and boat harness only.",
        size=9,
    )

    h(doc, "Three power changes (summary)", 1)
    bullets(
        doc,
        [
            "One Voltaic V50 USB-A #1 feed only for the Pico path: V50+ → INA219 Vin+ → TPS2113A IN1 → Pico VSYS (pin 39). "
            "INA219 measures total current on that path (Pico + anything you branch from the same 5 V bus upstream of the split).",
            "Remove a second V50 USB-A port dedicated only to modem power. Modem 5 V comes from the shared modem 5 V bus instead (see diodes below).",
            "Modem hardware reset: SIM7600 RST → Pico physical pin 14 (GP10). Common GND with Pico and modem.",
        ],
    )
    p(doc, "No extra Pico wire to modem PWR/PWRKEY if the HAT jumper is PWR–3V3 (auto power-on when 5 V is present).", bold=True, size=8.5)

    h(doc, "How power works: V50 + boat switch (TPS2113A + PlusRoc)", 1)
    p(doc, "Pololu TPS2113A picks the higher-priority 5 V source for the Pico:", size=9)
    bullets(
        doc,
        [
            "IN1 — V50 USB-A #1 → INA219 → TPS IN1 (battery switch OFF: Pico runs from V50).",
            "IN2 — PlusRoc 5 V OUT+ when master battery Switch is ON (charges V50 via USB-C and can feed Pico via TPS).",
            "OUT → Pico VSYS pin 39. GND common everywhere.",
        ],
    )
    p(doc, "PlusRoc 12 V → 5 V buck (waterproof): IN+ from boat Switch wire; GND to ground bus.", bold=True, size=9)
    tbl(doc, ["PlusRoc / node", "Connect to"], [
        ("IN+", "Boat Switch (same node as Opto Ch5 IN+ — harness wire #8)"),
        ("OUT+ 5 V", "TPS2113A IN2 (+)"),
        ("OUT+ 5 V", "V50 USB-C charge IN (+) direct — no diode on charge path"),
        ("OUT+ 5 V", "Modem 5 V bus through a diode (OR with V50-A#2)"),
        ("GND", "Ground bus (PlusRoc, V50 −, Pico, modem)"),
    ])

    h(doc, "V50 — three ports", 2)
    tbl(doc, ["Port", "Direction", "Connect to"], [
        ("USB-A #1", "OUT", "INA219 → TPS IN1 → VSYS (DONE)"),
        ("USB-A #2", "OUT", "(+) through diode → modem 5 V bus; (−) → ground"),
        ("USB-C", "IN (charge bank)", "(+) ← PlusRoc OUT+ direct; (−) → ground"),
    ])
    p(
        doc,
        "USB-C charges the V50 when Switch is ON (12 V on PlusRoc). Key wire senses ignition only — it does not feed PlusRoc. "
        "Do not tie USB-C into the Pico path or modem bus backwards.",
        size=8,
    )
    p(doc, "Passives on 5 V: diodes OR-ing modem feeds (PlusRoc OUT+ and V50 USB-A #2). RUN wake: 10 kΩ from opto OUT1–4 shared node to Pico pin 30. (No capacitor called out in as-built docs.)", size=8)

    h(doc, "SECTION 1 — Already done (do not change)", 1)
    tbl(doc, ["Pico pin", "GPIO", "Connected to"], [
        ("39", "VSYS", "TPS2113A OUT"),
        ("38", "—", "GND rail"),
        ("36", "3V3", "+ rail → sensor Vcc"),
        ("1", "GP0", "Engine INA260 SDA"),
        ("2", "GP1", "Engine INA260 SCL"),
        ("4", "GP2", "House INA260 SDA"),
        ("5", "GP3", "House INA260 SCL"),
        ("6", "GP4", "INA219 SDA"),
        ("7", "GP5", "INA219 SCL"),
        ("9", "GP6", "TPS STAT"),
        ("15", "GP11", "TPS VSNS (keep here)"),
        ("17", "GP13", "Red LED"),
        ("19", "GP14", "Blue LED"),
        ("20", "GP15", "Green LED"),
    ])
    p(doc, "INA260: VIN+ = battery/source side; VIN− = load side. Float 12 V tees to House VIN+.", size=8)
    p(doc, "Firmware VSNS: Pin(11, OUT) + value(0) — physical pin 15.", size=8)

    h(doc, "SECTION 2 — Still to do (new wires only)", 1)

    h(doc, "SIM7600X modem", 2)
    p(doc, "HAT jumpers: UART = B · PWR = PWR–3V3 · VCCIO = 3.3 V · Flight = NC", bold=True, size=8.5)
    tbl(doc, ["Modem", "Connect to"], [
        ("5 V", "Modem 5 V bus (diode-OR from PlusRoc and/or V50 USB-A #2) — not Pico 3V3"),
        ("GND", "GND rail"),
        ("RXD", "Pico pin 11 (GP8)"),
        ("TXD", "Pico pin 12 (GP9)"),
        ("RST", "Pico pin 14 (GP10) — reset when hung; not PWRKEY"),
        ("Antennas", "MAIN U.FL / GNSS SMA"),
    ])
    p(
        doc,
        "RST pulse recovers a stuck modem; it does not power on a module after AT+CPOF full power-off (that needs 5 V cycle or PWRKEY). Normal use: leave PWR on HAT jumper, use firmware reset only.",
        size=8,
    )

    h(doc, "Optocoupler — boat 12 V inputs", 2)
    tbl(doc, ["Opto IN", "Boat wire"], [
        ("Ch1 IN+", "Mid bilge"),
        ("Ch2 IN+", "Aft bilge"),
        ("Ch3 IN+", "Mid water (return)"),
        ("Ch4 IN+", "Aft water (return)"),
        ("Ch5 IN+", "Switch → also PlusRoc IN+"),
        ("Ch6 IN+", "Key"),
        ("All IN−", "Ground bus"),
    ])

    h(doc, "Optocoupler — Pico 3.3 V outputs", 2)
    tbl(doc, ["Opto OUT", "Pin", "GPIO", "Signal", "RUN wake?"], [
        ("OUT1", "31", "GP26", "Mid bilge pump", "Yes"),
        ("OUT2", "29", "GP22", "Aft bilge pump", "Yes"),
        ("OUT3", "25", "GP19", "Mid water float", "Yes"),
        ("OUT4", "24", "GP18", "Aft water float", "Yes"),
        ("OUT5", "26", "GP20", "Battery switch", "No"),
        ("OUT6", "27", "GP21", "Key", "No"),
        ("OUT7–8", "21, 22", "GP16–17", "Spare", "No"),
    ])
    p(doc, "Opto jumpers = pull-up. RUN: OUT1–4 → shared node → 10 kΩ → pin 30.", size=8)
    p(doc, "Each water float: one wire → House + (VIN+); other wire → Opto Ch3/4 IN+.", size=8)

    h(doc, "Switch vs key", 2)
    tbl(doc, ["Wire", "Meaning"], [
        ("Switch", "Master battery ON — powers PlusRoc and boat"),
        ("Key", "Ignition ON (sense on pin 27); does not power PlusRoc"),
    ])

    h(doc, "Field checklist", 2)
    tbl(doc, ["☐", "Task"], [
        ("☐", "Modem: 5 V, GND, pins 11, 12, 14; HAT jumpers set"),
        ("☐", "Modem 5 V diode bus from PlusRoc + V50 USB-A #2"),
        ("☐", "PlusRoc: Switch IN+, OUT+ → TPS IN2, V50 USB-C, modem (via diode)"),
        ("☐", "Opto outputs → pins 24–27, 29, 31 (+ spare 21–22)"),
        ("☐", "RUN: bilge/float OUT1–4 → 10 kΩ → pin 30"),
        ("☐", "16 boat harness wires (page 2)"),
    ], [0.25, 6.0])

    doc.add_page_break()

    h(doc, "SECTION 3 — Pin map (physical pin numbers)", 1)
    tbl(doc, ["Pin", "Status", "Goes to"], [
        ("1–7, 9, 15, 17–20, 39", "DONE", "I2C, TPS STAT/VSNS, LEDs, VSYS"),
        ("11, 12, 14", "TODO", "Modem RXD, TXD, RST"),
        ("24–27, 29, 31", "TODO", "Opto OUT4–OUT1, switch, key"),
        ("21–22", "Spare", "Opto Ch7–8"),
        ("30", "TODO", "RUN ← 10 kΩ ← opto OUT1–4"),
    ])

    h(doc, "SECTION 4 — 16 boat harness wires", 1)
    tbl(doc, ["#", "Label", "Connect to"], [
        ("1", "Engine +", "Bottom INA260 VIN+"),
        ("2", "Engine −", "Bottom GND + ground bus"),
        ("3", "House +", "Middle INA260 VIN+ + tee both float 12 V"),
        ("4", "House −", "Middle GND + ground bus"),
        ("5", "Mid bilge", "Opto Ch1 IN+"),
        ("6", "Aft bilge", "Opto Ch2 IN+"),
        ("7", "Key", "Opto Ch6 IN+"),
        ("8", "Switch", "PlusRoc IN+ + Opto Ch5 IN+"),
        ("9", "Mid water", "House + (wire 3)"),
        ("10", "Mid water", "Opto Ch3 IN+"),
        ("11", "Solar engine", "Bottom INA260 VIN+"),
        ("12", "Solar engine", "Bottom INA260 VIN−"),
        ("13", "Solar house", "Middle INA260 VIN+"),
        ("14", "Solar house", "Middle INA260 VIN−"),
        ("15", "Aft water", "House + (wire 3)"),
        ("16", "Aft water", "Opto Ch4 IN+"),
    ], [0.35, 1.05, 4.8])

    h(doc, "Firmware GPIO quick reference (config.py)", 1)
    tbl(doc, ["Function", "GPIO", "Physical pin"], [
        ("UART TX / RX", "8 / 9", "11 / 12"),
        ("Modem RST", "10", "14"),
        ("TPS STAT / VSNS", "6 / 11", "9 / 15"),
        ("Bilge mid / aft", "26 / 22", "31 / 29"),
        ("Float mid / aft", "19 / 18", "25 / 24"),
        ("Switch / Key", "20 / 21", "26 / 27"),
    ])

    p(doc, "Source: boat_monitor/WIRING_CHEAT_SHEET.md · Regenerate: python3 boat_monitor/build_cheat_sheet_docx.py", size=7)

    doc.save(OUT)
    print(f"Wrote {OUT}")


if __name__ == "__main__":
    main()
